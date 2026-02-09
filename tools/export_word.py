"""
Word Document Exporter

Converts markdown-formatted text back to a Word document (.docx) with proper formatting.
This allows users to export translated text with formatting preserved.
"""

import re
from io import BytesIO
from typing import List, Tuple, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX


# Map CSS color codes to Word highlight colors
CSS_TO_HIGHLIGHT = {
    "#FFFF00": WD_COLOR_INDEX.YELLOW,
    "#00FF00": WD_COLOR_INDEX.BRIGHT_GREEN,
    "#00FFFF": WD_COLOR_INDEX.TURQUOISE,
    "#FF00FF": WD_COLOR_INDEX.PINK,
    "#0000FF": WD_COLOR_INDEX.BLUE,
    "#FF0000": WD_COLOR_INDEX.RED,
    "#000080": WD_COLOR_INDEX.DARK_BLUE,
    "#008080": WD_COLOR_INDEX.TEAL,
    "#008000": WD_COLOR_INDEX.GREEN,
    "#800080": WD_COLOR_INDEX.VIOLET,
    "#800000": WD_COLOR_INDEX.DARK_RED,
    "#808000": WD_COLOR_INDEX.DARK_YELLOW,
    "#808080": WD_COLOR_INDEX.GRAY_50,
    "#C0C0C0": WD_COLOR_INDEX.GRAY_25,
}


def _hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """Convert hex color string to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def _strip_remaining_markers(text: str) -> str:
    """
    Strip any remaining unmatched markdown markers from text.
    This is a cleanup step for markers that weren't matched by the parser.
    """
    # Remove sequences of asterisks (bold/italic markers)
    text = re.sub(r'\*+', '', text)
    # Remove underline markers
    text = re.sub(r'\+\+', '', text)
    # Remove strikethrough markers
    text = re.sub(r'~~', '', text)
    # Remove highlight markers
    text = re.sub(r'==', '', text)
    # Remove color markers (::COLOR: or just ::)
    text = re.sub(r'::#[A-Fa-f0-9]{6}:', '', text)
    text = re.sub(r'::', '', text)
    return text


def _parse_formatted_text(text: str) -> List[dict]:
    """
    Parse markdown-formatted text into a list of text segments with formatting info.

    Each segment is a dict with:
    - text: The actual text content
    - bold: bool
    - italic: bool
    - underline: bool
    - strikethrough: bool
    - highlight: Optional[str] - hex color or None
    - color: Optional[str] - hex color or None
    """
    segments = []

    # Pattern to match all formatting markers
    # Order matters: process longer/more specific patterns first
    patterns = [
        # Colored highlight: ==#COLOR:text==
        (r'==(#[A-Fa-f0-9]{6}):(.+?)==', 'highlight_color'),
        # Simple highlight: ==text==
        (r'==(.+?)==', 'highlight'),
        # Colored text: ::COLOR:text::
        (r'::(#[A-Fa-f0-9]{6}):(.+?)::', 'color'),
        # Strikethrough: ~~text~~
        (r'~~(.+?)~~', 'strike'),
        # Underline: ++text++
        (r'\+\+(.+?)\+\+', 'underline'),
        # Bold and italic: ***text***
        (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),
        # Bold: **text**
        (r'\*\*(.+?)\*\*', 'bold'),
        # Italic: *text* (not preceded/followed by *)
        (r'(?<!\*)\*(.+?)\*(?!\*)', 'italic'),
    ]

    def parse_segment(text: str, inherited_format: dict = None) -> List[dict]:
        """Recursively parse text and extract formatted segments."""
        if inherited_format is None:
            inherited_format = {
                'bold': False,
                'italic': False,
                'underline': False,
                'strikethrough': False,
                'highlight': None,
                'color': None
            }

        result = []

        # Try each pattern
        for pattern, format_type in patterns:
            match = re.search(pattern, text, re.DOTALL)
            if match:
                # Get text before the match
                before = text[:match.start()]
                if before:
                    result.extend(parse_segment(before, inherited_format.copy()))

                # Create format for matched content
                new_format = inherited_format.copy()

                if format_type == 'highlight_color':
                    new_format['highlight'] = match.group(1)
                    inner_text = match.group(2)
                elif format_type == 'highlight':
                    new_format['highlight'] = '#FFFF00'  # Default yellow
                    inner_text = match.group(1)
                elif format_type == 'color':
                    new_format['color'] = match.group(1)
                    inner_text = match.group(2)
                elif format_type == 'strike':
                    new_format['strikethrough'] = True
                    inner_text = match.group(1)
                elif format_type == 'underline':
                    new_format['underline'] = True
                    inner_text = match.group(1)
                elif format_type == 'bold_italic':
                    new_format['bold'] = True
                    new_format['italic'] = True
                    inner_text = match.group(1)
                elif format_type == 'bold':
                    new_format['bold'] = True
                    inner_text = match.group(1)
                elif format_type == 'italic':
                    new_format['italic'] = True
                    inner_text = match.group(1)
                else:
                    inner_text = match.group(1)

                # Recursively parse inner text (for nested formatting)
                result.extend(parse_segment(inner_text, new_format))

                # Get text after the match
                after = text[match.end():]
                if after:
                    result.extend(parse_segment(after, inherited_format.copy()))

                return result

        # No patterns matched - return plain text segment
        if text:
            result.append({
                'text': text,
                **inherited_format
            })

        return result

    return parse_segment(text)


def _add_formatted_text_to_doc(doc: Document, text: str) -> None:
    """
    Add markdown-formatted text to a Word document.

    Args:
        doc: The Document object to add text to
        text: Markdown-formatted text
    """
    # Split text into paragraphs
    paragraphs = text.split('\n\n')

    for para_text in paragraphs:
        if not para_text.strip():
            doc.add_paragraph()
            continue

        # Handle single line breaks within paragraphs
        lines = para_text.split('\n')

        para = doc.add_paragraph()

        for i, line in enumerate(lines):
            # Parse the line into formatted segments
            segments = _parse_formatted_text(line)

            for segment in segments:
                # Clean any remaining unmatched markers from the text
                clean_text = _strip_remaining_markers(segment['text'])
                if not clean_text:
                    continue  # Skip empty segments
                run = para.add_run(clean_text)

                # Set font for this run (Times New Roman 12pt)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

                # Apply formatting
                if segment.get('bold'):
                    run.bold = True

                if segment.get('italic'):
                    run.italic = True

                if segment.get('underline'):
                    run.underline = True

                if segment.get('strikethrough'):
                    run.font.strike = True

                if segment.get('highlight'):
                    hex_color = segment['highlight'].upper()
                    # Try to use Word's built-in highlight colors
                    if hex_color in CSS_TO_HIGHLIGHT:
                        run.font.highlight_color = CSS_TO_HIGHLIGHT[hex_color]
                    else:
                        # For non-standard colors, use yellow as fallback
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                if segment.get('color'):
                    try:
                        r, g, b = _hex_to_rgb(segment['color'])
                        run.font.color.rgb = RGBColor(r, g, b)
                    except:
                        pass  # Keep default color if conversion fails

            # Add line break if not the last line
            if i < len(lines) - 1:
                para.add_run('\n')


def export_to_word(french_text: str, english_text: str = None, filename: str = "translation.docx") -> BytesIO:
    """
    Convert markdown-formatted text to a Word document.

    Creates a document with:
    1. French text (original) at the top
    2. A purple-highlighted separator line of asterisks
    3. English text (translated) below

    Args:
        french_text: Original French text with markdown formatting
        english_text: Translated English text with markdown formatting.
                      If None, only french_text is exported (backwards compatibility)
        filename: Name for the document (used in metadata)

    Returns:
        BytesIO object containing the Word document in Times New Roman 12pt
    """
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Add French text section
    _add_formatted_text_to_doc(doc, french_text)

    # If english_text is provided, add separator and English section
    if english_text is not None:
        # Add empty paragraph before separator
        doc.add_paragraph()

        # Add purple-highlighted asterisk separator
        separator_para = doc.add_paragraph()
        separator_run = separator_para.add_run('*' * 50)
        separator_run.font.name = 'Times New Roman'
        separator_run.font.size = Pt(12)
        separator_run.font.highlight_color = WD_COLOR_INDEX.VIOLET

        # Add empty paragraph after separator
        doc.add_paragraph()

        # Add English text section
        _add_formatted_text_to_doc(doc, english_text)

    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return output


def get_formatted_text_preview(text: str) -> str:
    """
    Get a preview of how the text will look when exported.
    Returns a description of the formatting found.
    """
    segments = _parse_formatted_text(text)

    formats_found = set()
    for seg in segments:
        if seg.get('bold'):
            formats_found.add('bold')
        if seg.get('italic'):
            formats_found.add('italic')
        if seg.get('underline'):
            formats_found.add('underline')
        if seg.get('strikethrough'):
            formats_found.add('strikethrough')
        if seg.get('highlight'):
            formats_found.add('highlight')
        if seg.get('color'):
            formats_found.add('colored text')

    if formats_found:
        return f"Formatting found: {', '.join(sorted(formats_found))}"
    else:
        return "No special formatting detected"


if __name__ == "__main__":
    # Test the export function
    test_text = """**Bold text** and *italic text* and ++underlined++.

This is a ==highlighted== word and ~~strikethrough~~.

This has ::FF0000:red text:: and ==#00FF00:green highlight==."""

    print("Test input:")
    print(test_text)
    print("\n" + "-" * 50)

    segments = _parse_formatted_text(test_text)
    print("\nParsed segments:")
    for seg in segments:
        formatting = []
        if seg.get('bold'):
            formatting.append('B')
        if seg.get('italic'):
            formatting.append('I')
        if seg.get('underline'):
            formatting.append('U')
        if seg.get('strikethrough'):
            formatting.append('S')
        if seg.get('highlight'):
            formatting.append(f'H:{seg["highlight"]}')
        if seg.get('color'):
            formatting.append(f'C:{seg["color"]}')

        format_str = ','.join(formatting) if formatting else 'plain'
        print(f"  [{format_str}] '{seg['text']}'")

    print("\n" + "-" * 50)
    print(get_formatted_text_preview(test_text))

    # Save test document
    output = export_to_word(test_text)
    with open("test_export.docx", "wb") as f:
        f.write(output.getvalue())
    print("\nTest document saved as 'test_export.docx'")
