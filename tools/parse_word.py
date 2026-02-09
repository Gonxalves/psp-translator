"""
Word Document Parser

Extracts text from Word documents (.docx) while preserving formatting
as markdown (bold, italic, line breaks, etc.).
"""

from io import BytesIO
from typing import Union
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX


# Map Word highlight colors to CSS colors
HIGHLIGHT_COLORS = {
    WD_COLOR_INDEX.YELLOW: "#FFFF00",
    WD_COLOR_INDEX.BRIGHT_GREEN: "#00FF00",
    WD_COLOR_INDEX.TURQUOISE: "#00FFFF",
    WD_COLOR_INDEX.PINK: "#FF00FF",
    WD_COLOR_INDEX.BLUE: "#0000FF",
    WD_COLOR_INDEX.RED: "#FF0000",
    WD_COLOR_INDEX.DARK_BLUE: "#000080",
    WD_COLOR_INDEX.TEAL: "#008080",
    WD_COLOR_INDEX.GREEN: "#008000",
    WD_COLOR_INDEX.VIOLET: "#800080",
    WD_COLOR_INDEX.DARK_RED: "#800000",
    WD_COLOR_INDEX.DARK_YELLOW: "#808000",
    WD_COLOR_INDEX.GRAY_50: "#808080",
    WD_COLOR_INDEX.GRAY_25: "#C0C0C0",
    WD_COLOR_INDEX.BLACK: "#000000",
}


def _get_highlight_color(run) -> str:
    """Get the highlight color from a run, if any."""
    try:
        if run.font.highlight_color and run.font.highlight_color != WD_COLOR_INDEX.AUTO:
            return HIGHLIGHT_COLORS.get(run.font.highlight_color, "#FFFF00")
    except:
        pass
    return None


def _get_font_color(run) -> str:
    """Get the font color from a run, if any."""
    try:
        if run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            return f"#{rgb}"
    except:
        pass
    return None


def parse_word_document(file: Union[str, BytesIO]) -> str:
    """
    Parse a Word document and extract text with formatting preserved as markdown.

    Args:
        file: Either a file path string or a BytesIO object (from Streamlit file_uploader)

    Returns:
        Text content with markdown formatting:
        - **bold** for bold text
        - *italic* for italic text
        - ++underline++ for underlined text
        - ~~strikethrough~~ for strikethrough text
        - ==highlighted== or ==#COLOR:highlighted== for highlighted text
        - ::COLOR:colored text:: for colored text
        - Proper line breaks and paragraph spacing

    Raises:
        Exception: If the document cannot be parsed
    """
    try:
        doc = Document(file)
        result_paragraphs = []

        for para in doc.paragraphs:
            # Skip empty paragraphs
            if not para.text.strip():
                result_paragraphs.append("")
                continue

            # Process each run (text segment with consistent formatting) in the paragraph
            paragraph_text = ""

            for run in para.runs:
                text = run.text

                if not text:
                    continue

                # Get additional formatting
                is_underline = run.underline is not None and run.underline != False
                is_strike = run.font.strike if run.font.strike else False
                highlight_color = _get_highlight_color(run)
                font_color = _get_font_color(run)

                # Apply basic formatting as markdown
                if run.bold and run.italic:
                    text = f"***{text}***"
                elif run.bold:
                    text = f"**{text}**"
                elif run.italic:
                    text = f"*{text}*"

                # Apply underline
                if is_underline:
                    text = f"++{text}++"

                # Apply strikethrough
                if is_strike:
                    text = f"~~{text}~~"

                # Apply highlight (with color if not yellow)
                if highlight_color:
                    if highlight_color == "#FFFF00":
                        text = f"=={text}=="
                    else:
                        text = f"=={highlight_color}:{text}=="

                # Apply font color (skip black as it's default)
                if font_color and font_color.upper() != "#000000":
                    text = f"::{font_color}:{text}::"

                paragraph_text += text

            # Clean up formatting markers that might have been split across runs
            paragraph_text = _clean_formatting(paragraph_text)

            result_paragraphs.append(paragraph_text)

        # Join paragraphs with double newlines (standard markdown paragraph separator)
        result = "\n\n".join(result_paragraphs)

        # Clean up excessive whitespace
        result = _clean_whitespace(result)

        return result

    except Exception as e:
        raise Exception(f"Failed to parse Word document: {e}")


def _clean_formatting(text: str) -> str:
    """
    Clean up formatting markers that might have been split across Word runs.

    When Word splits formatted text into multiple runs, we may get adjacent markers
    like **text****more** which should become **textmore**.
    """
    import re

    # Simplify asterisk sequences by removing adjacent end+start markers
    # 6+ asterisks → reduce to manageable (likely malformed)
    text = re.sub(r'\*{6,}', '***', text)
    # 5 asterisks (bold end + bold+italic start) → 3 asterisks
    text = re.sub(r'\*{5}', '***', text)
    # 4 asterisks (bold end + bold start) → nothing (merge)
    text = re.sub(r'\*{4}', '', text)

    # Merge adjacent underline markers: ++text++++more++ → ++textmore++
    text = re.sub(r'\+{4,}', '', text)

    # Merge adjacent strikethrough markers: ~~text~~~~more~~ → ~~textmore~~
    text = re.sub(r'~{4,}', '', text)

    # Merge adjacent highlight markers: ==text====more== → ==textmore==
    text = re.sub(r'={4,}', '', text)

    # Merge adjacent color markers: ::color:text::::color:more:: → ::color:textmore::
    text = re.sub(r':{4,}', '', text)

    return text


def _clean_whitespace(text: str) -> str:
    """
    Clean up excessive whitespace while preserving intentional line breaks.
    """
    import re

    # Replace 3+ newlines with 2 (paragraph break)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove trailing whitespace from lines
    text = re.sub(r'[ \t]+$', '', text, flags=re.MULTILINE)

    # Remove leading/trailing whitespace from entire text
    text = text.strip()

    return text


def get_document_info(file: Union[str, BytesIO]) -> dict:
    """
    Get information about a Word document.

    Args:
        file: Either a file path string or a BytesIO object

    Returns:
        Dictionary with document info:
        - paragraph_count: Number of paragraphs
        - word_count: Approximate word count
        - character_count: Character count
    """
    try:
        doc = Document(file)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = " ".join(paragraphs)

        return {
            'paragraph_count': len(paragraphs),
            'word_count': len(full_text.split()),
            'character_count': len(full_text)
        }

    except Exception as e:
        return {
            'error': str(e),
            'paragraph_count': 0,
            'word_count': 0,
            'character_count': 0
        }


if __name__ == "__main__":
    # Test with a sample document
    import sys

    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        print(f"Parsing: {file_path}")
        print("-" * 50)

        try:
            text = parse_word_document(file_path)
            print(text)
            print("-" * 50)

            info = get_document_info(file_path)
            print(f"Document info: {info}")

        except Exception as e:
            print(f"Error: {e}")
    else:
        print("Usage: python parse_word.py <path_to_docx>")
